import os
import re
from PyQt6.QtCore import Qt, QMarginsF
from PyQt6.QtGui import (
    QTextDocument, QTextCursor, QPageLayout, QPageSize,
    QFont, QColor
)
from PyQt6.QtPrintSupport import QPrinter

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

try:
    import markdown
except ImportError:
    markdown = None

class DocumentManager:
    @staticmethod
    def _is_docx_file(filepath):
        try:
            with open(filepath, "rb") as f:
                return f.read(4) == b"PK\x03\x04"
        except Exception:
            return False

    @staticmethod
    def load_file(filepath, text_document: QTextDocument):
        """Loads a file (.docx, .md, .html, .txt) into a QTextDocument preserving rich formatting."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".docx" or (not ext and DocumentManager._is_docx_file(filepath)):
            if docx is None:
                raise ImportError("python-docx is required to read .docx files.")
            doc = docx.Document(filepath)
            html_parts = []
            
            for p in doc.paragraphs:
                p_text = p.text
                if not p_text.strip() and not p.runs:
                    html_parts.append("<p><br/></p>")
                    continue

                style_name = p.style.name.lower() if p.style else ""
                tag = "p"
                if "heading 1" in style_name:
                    tag = "h1"
                elif "heading 2" in style_name:
                    tag = "h2"
                elif "heading 3" in style_name:
                    tag = "h3"
                elif "title" in style_name:
                    tag = "h1"
                elif "quote" in style_name:
                    tag = "blockquote"

                align_style = ""
                if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align_style = "text-align: center;"
                elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align_style = "text-align: right;"
                elif p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    align_style = "text-align: justify;"

                runs_html = []
                for run in p.runs:
                    txt = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
                    
                    span_styles = []
                    if run.font.name:
                        span_styles.append(f"font-family: '{run.font.name}';")
                    if run.font.size:
                        span_styles.append(f"font-size: {run.font.size.pt}pt;")
                    if run.font.color and run.font.color.rgb:
                        span_styles.append(f"color: #{run.font.color.rgb};")
                        
                    if run.bold:
                        txt = f"<b>{txt}</b>"
                    if run.italic:
                        txt = f"<i>{txt}</i>"
                    if run.underline:
                        txt = f"<u>{txt}</u>"
                    if run.font.strike:
                        txt = f"<s>{txt}</s>"
                        
                    if span_styles:
                        txt = f'<span style="{" ".join(span_styles)}">{txt}</span>'
                        
                    runs_html.append(txt)

                inner = "".join(runs_html) if runs_html else p_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                style_attr = f' style="{align_style}"' if align_style else ""
                html_parts.append(f"<{tag}{style_attr}>{inner}</{tag}>")

            # Load docx tables if present
            for tbl in doc.tables:
                tbl_html = ['<table border="1" cellpadding="6" style="border-collapse: collapse; margin: 12px 0;">']
                for row in tbl.rows:
                    tbl_html.append("<tr>")
                    for cell in row.cells:
                        cell_txt = cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        tbl_html.append(f"<td>{cell_txt}</td>")
                    tbl_html.append("</tr>")
                tbl_html.append("</table>")
                html_parts.append("".join(tbl_html))

            text_document.setHtml("\n".join(html_parts))

        elif ext in {".md", ".markdown"}:
            with open(filepath, "r", encoding="utf-8") as f:
                raw_md = f.read()
            if markdown:
                html = markdown.markdown(raw_md, extensions=["tables", "fenced_code", "nl2br", "sane_lists"])
                text_document.setHtml(html)
            else:
                text_document.setPlainText(raw_md)

        elif ext in {".html", ".htm"}:
            with open(filepath, "r", encoding="utf-8") as f:
                html = f.read()
            text_document.setHtml(html)

        else:  # .txt or fallback
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
            text_document.setPlainText(text)

    @staticmethod
    def save_file(filepath, text_document: QTextDocument):
        """Saves a QTextDocument to the specified format (.docx, .pdf, .md, .html, .txt) with full formatting."""
        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".docx" or not ext:
            if docx is None:
                raise ImportError("python-docx is required to write .docx files.")
            
            doc = docx.Document()
            
            # Standard A4 margins (20mm)
            for section in doc.sections:
                section.top_margin = Inches(0.79)
                section.bottom_margin = Inches(0.79)
                section.left_margin = Inches(0.79)
                section.right_margin = Inches(0.79)
                
            block = text_document.begin()
            while block.isValid():
                fmt = block.blockFormat()
                heading_level = fmt.headingLevel()
                alignment = fmt.alignment()
                text_list = block.textList()
                
                # Check heading / style
                if heading_level == 1:
                    p = doc.add_heading(level=1)
                elif heading_level == 2:
                    p = doc.add_heading(level=2)
                elif heading_level == 3:
                    p = doc.add_heading(level=3)
                elif text_list is not None:
                    p = doc.add_paragraph(style="List Bullet")
                else:
                    p = doc.add_paragraph()
                    
                # Alignment
                if alignment & Qt.AlignmentFlag.AlignHCenter:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment & Qt.AlignmentFlag.AlignRight:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment & Qt.AlignmentFlag.AlignJustify:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif alignment & Qt.AlignmentFlag.AlignLeft:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                # Walk character fragments in block
                it = block.begin()
                while not it.atEnd():
                    frag = it.fragment()
                    if frag.isValid():
                        raw_txt = frag.text()
                        # Filter Qt object replacement chars (e.g. embedded tables/images)
                        txt = raw_txt.replace('\ufffc', '')
                        if txt:
                            char_fmt = frag.charFormat()
                            run = p.add_run(txt)
                            
                            # Bold
                            if char_fmt.fontWeight() >= 600 or char_fmt.font().bold():
                                run.bold = True
                            # Italic
                            if char_fmt.fontItalic():
                                run.italic = True
                            # Underline
                            if char_fmt.fontUnderline():
                                run.underline = True
                            # Strikethrough
                            if char_fmt.fontStrikeOut():
                                run.font.strike = True
                            # Font Family
                            try:
                                fam = char_fmt.font().family()
                                if fam and fam.lower() not in {"default", "sans-serif", "serif"}:
                                    run.font.name = fam
                            except Exception:
                                pass
                            # Font Size
                            pt_sz = char_fmt.fontPointSize()
                            if pt_sz > 0:
                                run.font.size = Pt(pt_sz)
                            # Text Color
                            fg = char_fmt.foreground().color()
                            if fg.isValid() and fg.name() != "#000000":
                                run.font.color.rgb = RGBColor(fg.red(), fg.green(), fg.blue())
                                
                    it += 1
                    
                block = block.next()
                
            doc.save(filepath)

        elif ext == ".pdf":
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filepath)
            
            # A4 page layout with 20mm standard margins
            page_layout = QPageLayout(
                QPageSize(QPageSize.PageSizeId.A4),
                QPageLayout.Orientation.Portrait,
                QMarginsF(20, 20, 20, 20),
                QPageLayout.Unit.Millimeter
            )
            printer.setPageLayout(page_layout)
            text_document.print(printer)

        elif ext in {".html", ".htm"}:
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(text_document.toHtml())

        elif ext in {".md", ".markdown"}:
            md_text = DocumentManager.document_to_markdown(text_document)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(md_text)

        else:  # .txt or plain text
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(text_document.toPlainText())

    @staticmethod
    def document_to_markdown(text_document: QTextDocument) -> str:
        """Converts a QTextDocument to rich Markdown preserving headings, bold, italic, strike, lists, quotes."""
        md_lines = []
        block = text_document.begin()
        
        while block.isValid():
            fmt = block.blockFormat()
            heading_level = fmt.headingLevel()
            text_list = block.textList()
            is_quote = fmt.leftMargin() >= 20
            
            # Prefix for block
            prefix = ""
            if heading_level == 1:
                prefix = "# "
            elif heading_level == 2:
                prefix = "## "
            elif heading_level == 3:
                prefix = "### "
            elif text_list is not None:
                prefix = "* "
            elif is_quote:
                prefix = "> "
                
            line_parts = []
            it = block.begin()
            while not it.atEnd():
                frag = it.fragment()
                if frag.isValid():
                    txt = frag.text().replace('\ufffc', '')
                    if txt:
                        cf = frag.charFormat()
                        is_bold = cf.fontWeight() >= 600 or cf.font().bold()
                        is_italic = cf.fontItalic()
                        is_strike = cf.fontStrikeOut()
                        is_underline = cf.fontUnderline()
                        
                        part = txt
                        if is_bold and is_italic:
                            part = f"***{part}***"
                        elif is_bold:
                            part = f"**{part}**"
                        elif is_italic:
                            part = f"*{part}*"
                        elif is_underline:
                            part = f"_{part}_"
                        if is_strike:
                            part = f"~~{part}~~"
                        line_parts.append(part)
                it += 1
                
            line_str = "".join(line_parts)
            if prefix and line_str:
                md_lines.append(f"{prefix}{line_str}\n")
            elif line_str:
                md_lines.append(f"{line_str}\n")
            else:
                md_lines.append("\n")
                
            block = block.next()
            
        full_md = "\n".join(md_lines)
        # Collapse 3+ consecutive newlines
        full_md = re.sub(r'\n{3,}', '\n\n', full_md)
        return full_md.strip() + "\n"
